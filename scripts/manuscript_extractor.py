#!/usr/bin/env python3
"""
Manuscript Statistical Extractor v2
Extract reported statistics from manuscript text (DOCX, PDF, or pasted text).
Fixes: decision-flip detection, separate p=/p</p>/ns handling,
correct r(df) formula, expanded coverage (negatives, Unicode, LaTeX, LLM metrics).
"""

import re
import json
import argparse
import sys
from pathlib import Path
from dataclasses import dataclass, asdict
from typing import List, Optional, Dict, Any

try:
    from scipy import stats as sp_stats
    HAS_SCIPY = True
except ImportError:
    HAS_SCIPY = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    HAS_PDFMINER = True
except ImportError:
    try:
        import fitz
        HAS_PDFMINER = False
        HAS_FITZ = True
    except ImportError:
        HAS_PDFMINER = False
        HAS_FITZ = False


@dataclass
class NHSTResult:
    location: str
    raw_text: str
    test_type: str
    test_statistic: float
    df1: Optional[float]
    df2: Optional[float]
    reported_p_expr: str
    reported_p_value: Optional[float]
    sign: Optional[str]
    computed_p: Optional[float] = None
    consistent: Optional[bool] = None
    decision_error: bool = False
    decision_reported: Optional[str] = None
    decision_computed: Optional[str] = None
    check_type: str = "deterministic"


@dataclass
class DescriptiveResult:
    location: str
    raw_text: str
    variable: Optional[str]
    mean: Optional[float]
    sd: Optional[float]
    sem: Optional[float]
    n: Optional[int]
    spread_label: Optional[str]
    scale_type: str = "unknown"
    scale_min: Optional[float] = None
    scale_max: Optional[float] = None
    ci_lower: Optional[float] = None
    ci_upper: Optional[float] = None


@dataclass
class ProportionResult:
    location: str
    raw_text: str
    proportion: float
    count: Optional[int]
    total_n: int
    reported_sd: Optional[float] = None


@dataclass
class LLMMetricResult:
    location: str
    raw_text: str
    metric: str
    value: float
    value_normalized: Optional[float] = None  # 0-1 range
    unit: str = "ratio"  # "ratio" | "percent" | "ms" | "tokens" | "dollars"
    ci_lower: Optional[float] = None
    ci_upper: Optional[float] = None
    std: Optional[float] = None
    seeds: Optional[int] = None


# Numeric token: supports negatives, decimals without leading 0
NUM = r"-?(?:\d+\.?\d*|\.\d+)"

PATTERNS = {
    "t": re.compile(
        r"(?:^|[^A-Za-z])t\s*\(\s*(" + NUM + r")\s*\)\s*=?\s*(" + NUM + r")\s*[,;]?\s*"
        r"(?:p\s*([<>=])\s*(" + NUM + r")|p\s*=\s*(" + NUM + r")|ns)",
        re.IGNORECASE,
    ),
    "F": re.compile(
        r"(?:^|[^A-Za-z])F\s*\(\s*(" + NUM + r")\s*,\s*(" + NUM + r")\s*\)\s*=?\s*(" + NUM + r")\s*[,;]?\s*"
        r"(?:p\s*([<>=])\s*(" + NUM + r")|p\s*=\s*(" + NUM + r")|ns)",
        re.IGNORECASE,
    ),
    "chi2": re.compile(
        r"(?:\\chi\^?2|\\chi\s*\^?\s*2|chi[-\s]?square|chi2?|χ\s*[²2])\s*\(\s*(" + NUM + r")\s*\)\s*=?\s*(" + NUM + r")\s*[,;]?\s*"
        r"(?:p\s*([<>=])\s*(" + NUM + r")|p\s*=\s*(" + NUM + r")|ns)",
        re.IGNORECASE,
    ),
    "r": re.compile(
        r"(?:^|[^A-Za-z])r\s*(?:s|pb|sb|sp)?\s*\(\s*(" + NUM + r")\s*\)\s*=?\s*(" + NUM + r")\s*[,;]?\s*"
        r"(?:p\s*([<>=])\s*(" + NUM + r")|p\s*=\s*(" + NUM + r")|ns)",
        re.IGNORECASE,
    ),
    "z": re.compile(
        r"(?:^|[^A-Za-z])z\s*=?\s*(" + NUM + r")\s*[,;]?\s*"
        r"(?:p\s*([<>=])\s*(" + NUM + r")|p\s*=\s*(" + NUM + r")|ns)",
        re.IGNORECASE,
    ),
    "Q": re.compile(
        r"(?:^|[^A-Za-z])Q\s*\(\s*(" + NUM + r")\s*\)\s*=?\s*(" + NUM + r")\s*[,;]?\s*"
        r"(?:p\s*([<>=])\s*(" + NUM + r")|p\s*=\s*(" + NUM + r")|ns)",
        re.IGNORECASE,
    ),
    "W": re.compile(
        r"(?:^|[^A-Za-z])(?:W|U)\s*\(\s*(" + NUM + r")\s*\)\s*=?\s*(" + NUM + r")\s*[,;]?\s*"
        r"(?:p\s*([<>=])\s*(" + NUM + r")|p\s*=\s*(" + NUM + r")|ns)",
        re.IGNORECASE,
    ),
}

DESCRIPTIVE_PATTERNS = [
    re.compile(
        r"(?:M|mean)\s*=?\s*(" + NUM + r")\s*(?:±|\\pm|\+\s*[-–—])\s*(" + NUM + r")"
        r"(?:\s*\((SD|SEM|SE|sd|sem|se)\))?"
        r"(?:\s*[,;]?\s*(?:n|N)\s*=\s*(\d+))?",
        re.IGNORECASE,
    ),
    re.compile(
        r"(?:M|mean)\s*=?\s*(" + NUM + r")\s*[,;]?\s*(?:SD|sd|SEM|sem)\s*=?\s*(" + NUM + r")"
        r"(?:\s*[,;]?\s*(?:n|N)\s*=\s*(\d+))?",
        re.IGNORECASE,
    ),
]

PROPORTION_PATTERNS = [
    re.compile(r"(" + NUM + r")\s*%\s*\((\d+)\s*/\s*(\d+)\)"),
    re.compile(r"(\d+)\s*/\s*(\d+)\s*\((" + NUM + r")%\)"),
]

LLM_PATTERNS = [
    # pass@k variants: pass@1 = 0.45, pass@10: .75, pass @ 5 = 0.3
    re.compile(r"pass\s*@\s*(\d+)\s*[:=]\s*(" + NUM + r")", re.IGNORECASE),
    # accuracy / F1 / EM / BLEU / ROUGE
    re.compile(r"(?:accuracy|acc|F1|BLEU|ROUGE|exact match|EM|AUC|AUROC|AUPRC)\s*[:=]\s*(" + NUM + r")\s*(?:%|percent)?", re.IGNORECASE),
    # win rate / success rate
    re.compile(r"(?:win rate|win_rate|success rate|success_rate)\s*[:=]\s*(" + NUM + r")\s*(?:%|percent)?", re.IGNORECASE),
    # cost / latency / tokens / tool calls
    re.compile(r"(?:latency|token[s]?|API call[s]?|tool call[s]?|cost)\s*[:=]\s*(" + NUM + r")", re.IGNORECASE),
    # std over seeds: mean ± std over 5 seeds
    re.compile(r"mean\s*±\s*(?:std|sd|standard deviation)\s*over\s*(\d+)\s*seed[s]?", re.IGNORECASE),
    # mean ± std: (already covered by DESCRIPTIVE_PATTERNS but keep for explicit LLM context)
    re.compile(r"(?:mean\s*±\s*std|mean\s*±\s*sd)\s*[:=]?\s*(" + NUM + r")\s*[,;]?\s*(" + NUM + r")", re.IGNORECASE),
]


# ---------------------------------------------------------------------------
# Core functions
# ---------------------------------------------------------------------------

def strip_punctuation(s: str) -> str:
    if not s:
        return s
    return s.rstrip(".,;:) ")


def parse_p_value(sign_str: str, value_str: str) -> tuple:
    if not sign_str and not value_str:
        return None, None
    sign = sign_str if sign_str else "="
    if not value_str:
        return sign, None
    value_str = strip_punctuation(value_str)
    try:
        return sign, float(value_str)
    except (ValueError, TypeError):
        return sign, None


def compute_p_value(test_type: str, statistic: float, df1: float,
                    df2: Optional[float] = None) -> Optional[float]:
    if not HAS_SCIPY:
        return None
    try:
        s = float(statistic)
        if test_type == "t":
            return float(sp_stats.t.sf(abs(s), float(df1)) * 2)
        elif test_type == "F":
            if df2 is None:
                return None
            return float(sp_stats.f.sf(s, float(df1), float(df2)))
        elif test_type == "chi2":
            return float(sp_stats.chi2.sf(s, float(df1)))
        elif test_type == "r":
            if abs(s) >= 1:
                return None
            df = float(df1)
            if df <= 0:
                return None
            t_stat = s * (df / (1 - s**2)) ** 0.5
            return float(sp_stats.t.sf(abs(t_stat), df) * 2)
        elif test_type in ("z", "Q"):
            return float(sp_stats.norm.sf(abs(s)) * 2)
        elif test_type in ("W",):
            return float(sp_stats.norm.sf(abs(s)) * 2)
    except Exception:
        return None
    return None


def decide_significance(p_value: Optional[float], alpha: float = 0.05) -> str:
    if p_value is None:
        return "unknown"
    return "significant" if p_value < alpha else "nonsignificant"


def reported_decision(sign: Optional[str], reported_val: Optional[float],
                      ns: bool = False, alpha: float = 0.05) -> str:
    """Determine what the author CLAIMED (not what the math says)."""
    if ns:
        return "nonsignificant"
    if sign == "<" and reported_val is not None:
        # p < .05 means author claims significance at .05
        return "significant" if reported_val <= alpha else "unknown"
    if sign == ">" and reported_val is not None:
        return "nonsignificant" if reported_val >= alpha else "unknown"
    if sign == "=" and reported_val is not None:
        return "significant" if reported_val < alpha else "nonsignificant"
    return "unknown"


def check_p_consistency(computed: Optional[float], sign: Optional[str],
                        reported_val: Optional[float], ns: bool = False) -> tuple:
    if computed is None:
        return None, False

    # Determine reported decision
    if ns:
        reported_sig = False
    elif sign == "<" and reported_val is not None:
        reported_sig = reported_val <= 0.05
    elif sign == ">" and reported_val is not None:
        reported_sig = False
    elif sign == "=" and reported_val is not None:
        reported_sig = reported_val < 0.05
    else:
        reported_sig = None

    computed_sig = computed < 0.05

    # Priority 1: decision flip is ALWAYS flagged
    if reported_sig is not None and computed_sig != reported_sig:
        return False, True

    # Priority 2: numerical consistency for exact p= values
    if sign == "=" and reported_val is not None:
        return abs(computed - reported_val) < 0.01, False

    if sign == "<" and reported_val is not None:
        return (computed < reported_val) or abs(computed - reported_val) < 0.01, False

    if sign == ">" and reported_val is not None:
        return (computed > reported_val) or abs(computed - reported_val) < 0.01, False

    return None, False


def _extract_scale_bounds(text: str) -> tuple:
    """Try to extract (min, max) from patterns like '1-7 scale', '0 to 10'."""
    # Pattern: 1-7 scale, 1–7 Likert, 0-10 rating
    m = re.search(r"(\d+)\s*[-–—]\s*(\d+)\s*(?:scale|likert|rating|point|item)?", text.lower())
    if m:
        return int(m.group(1)), int(m.group(2))
    # Pattern: 1 to 7
    m = re.search(r"(\d+)\s+to\s+(\d+)", text.lower())
    if m:
        return int(m.group(1)), int(m.group(2))
    return None, None


def infer_scale_type(text_context: str) -> Dict[str, Any]:
    """
    Infer scale type from LOCAL context (same sentence/phrase).
    Returns dict with scale_type, confidence, evidence, scale_min, scale_max.
    """
    lower = text_context.lower()
    scale_min, scale_max = _extract_scale_bounds(text_context)

    # Sentence-level continuous markers (HIGH confidence)
    strong_continuous = [
        "reaction time", "response time", "latency", "milliseconds",
        "age", "years old", "body mass", "bmi", "weight", "height",
        "temperature", "celsius", "fahrenheit",
        "blood pressure", "heart rate", "concentration",
        "0-100 scale", "0–100 scale", "vas", "visual analog",
        "continuous scale", "continuous variable",
    ]
    for m in strong_continuous:
        if re.search(r'\b' + re.escape(m) + r'\b', lower):
            return {"scale_type": "continuous", "confidence": "high",
                    "evidence": f"local context: '{m}'"}

    # Sentence-level integer markers (HIGH confidence)
    strong_integer = [
        "likert", "1-5 scale", "1–5 scale", "1-7 scale", "1–7 scale",
        "1 to 5", "1 to 7", "0-10 scale", "0–10 scale",
        "integer score", "discrete score", "ordinal scale",
    ]
    for m in strong_integer:
        if re.search(r'\b' + re.escape(m) + r'\b', lower):
            return {"scale_type": "integer", "confidence": "high",
                    "evidence": f"local context: '{m}'",
                    "scale_min": scale_min or 1, "scale_max": scale_max or 5}

    # Weaker continuous markers (MEDIUM confidence)
    weak_continuous = [
        "continuous", "mm", "cm", "kg", "ml", "seconds", "minutes",
        "percent", "percentage",
    ]
    for m in weak_continuous:
        if re.search(r'\b' + re.escape(m) + r'\b', lower):
            return {"scale_type": "continuous", "confidence": "medium",
                    "evidence": f"weak marker: '{m}'"}

    # Weaker integer markers (MEDIUM confidence)
    weak_integer = [
        "count", "frequency", "number of", "score", "rating",
        "0-10", "0–10", "1-5", "1–5", "1-7", "1–7",
    ]
    for m in weak_integer:
        if re.search(r'\b' + re.escape(m) + r'\b', lower):
            return {"scale_type": "integer", "confidence": "medium",
                    "evidence": f"weak marker: '{m}'",
                    "scale_min": scale_min or 1, "scale_max": scale_max or 5}

    # Scale bounds mention without explicit type marker
    if scale_min is not None and scale_max is not None:
        if scale_max - scale_min <= 10:
            return {"scale_type": "integer", "confidence": "medium",
                    "evidence": f"bounded scale {scale_min}-{scale_max}",
                    "scale_min": scale_min, "scale_max": scale_max}

    return {"scale_type": "unknown", "confidence": "low",
            "evidence": "no local scale indicator found"}


def extract_nhst(text: str, location: str = "text") -> List[NHSTResult]:
    results = []
    for test_type, pattern in PATTERNS.items():
        for match in pattern.finditer(text):
            try:
                groups = match.groups()
                raw = match.group(0)
                ns = groups[-1] is not None and groups[-1].lower() == "ns"

                if test_type == "F":
                    df1 = float(groups[0])
                    df2 = float(groups[1])
                    stat = float(groups[2])
                    sign, p_val = parse_p_value(groups[3], groups[4])
                    if not p_val and groups[5]:
                        sign, p_val = "=", float(strip_punctuation(groups[5]))
                elif test_type in ("t", "chi2", "r", "Q"):
                    df1 = float(groups[0])
                    df2 = None
                    stat = float(groups[1])
                    sign, p_val = parse_p_value(groups[2], groups[3])
                    if not p_val and groups[4]:
                        sign, p_val = "=", float(strip_punctuation(groups[4]))
                elif test_type == "z":
                    df1 = float("inf")
                    df2 = None
                    stat = float(groups[0])
                    sign, p_val = parse_p_value(groups[1], groups[2])
                    if not p_val and groups[3]:
                        sign, p_val = "=", float(strip_punctuation(groups[3]))
                elif test_type == "W":
                    df1 = float(groups[0])
                    df2 = None
                    stat = float(groups[1])
                    sign, p_val = parse_p_value(groups[2], groups[3])
                    if not p_val and groups[4]:
                        sign, p_val = "=", float(strip_punctuation(groups[4]))
                else:
                    continue

                computed = compute_p_value(test_type, stat, df1, df2)
                consistent, decision_error = check_p_consistency(
                    computed, sign, p_val, ns=ns
                )

                # P0-2 fix: use reported_decision() not decide_significance(alpha)
                # p<.05 should yield "significant", not "nonsignificant"
                decision_reported = reported_decision(sign, p_val, ns=ns)
                decision_computed = decide_significance(computed)

                results.append(NHSTResult(
                    location=location,
                    raw_text=raw.strip(),
                    test_type=test_type,
                    test_statistic=stat,
                    df1=df1,
                    df2=df2,
                    reported_p_expr=f"{sign} {p_val}" if sign and p_val else ("ns" if ns else "?"),
                    reported_p_value=p_val,
                    sign=sign,
                    computed_p=computed,
                    consistent=consistent,
                    decision_error=decision_error,
                    decision_reported=decision_reported if not ns else "nonsignificant",
                    decision_computed=decision_computed,
                ))
            except (ValueError, IndexError, TypeError):
                continue
    return results


def extract_descriptives(text: str, location: str = "text") -> List[DescriptiveResult]:
    results = []
    for pattern in DESCRIPTIVE_PATTERNS:
        for match in pattern.finditer(text):
            try:
                groups = match.groups()
                raw = match.group(0)
                mean_val = float(groups[0])
                spread_val = float(groups[1])
                label = "SD"
                n_val = None

                # Pattern differentiation: pat1 has label as group 2, n as group 3
                # pat2 has n as group 2 (no label group)
                if len(groups) >= 4:
                    # Pattern 1: (mean, spread, label, n)
                    if groups[2]:
                        label = groups[2].upper()
                    if groups[3]:
                        n_val = int(groups[3])
                elif len(groups) == 3:
                    # Pattern 2: (mean, spread, n)
                    if groups[2]:
                        n_val = int(groups[2])

                sd_val = spread_val if label in ("SD",) else None
                sem_val = spread_val if label in ("SEM", "SE") else None

                # P1-1: Use SENTENCE-LOCAL context only (same sentence/clause)
                # Find sentence boundaries around match
                sent_start = text.rfind(".", 0, match.start()) + 1
                sent_end = text.find(".", match.end())
                if sent_end == -1:
                    sent_end = len(text)
                ctx = text[sent_start:sent_end].strip()
                scale_info = infer_scale_type(ctx)
                scale_type = scale_info["scale_type"]
                scale_conf = scale_info.get("confidence", "low")
                scale_min = scale_info.get("scale_min")
                scale_max = scale_info.get("scale_max")

                results.append(DescriptiveResult(
                    location=location,
                    raw_text=raw.strip(),
                    variable=None,
                    mean=mean_val,
                    sd=sd_val,
                    sem=sem_val,
                    n=n_val,
                    spread_label=label,
                    scale_type=scale_type,
                    scale_min=scale_min,
                    scale_max=scale_max,
                ))
            except (ValueError, TypeError):
                continue
    return results


def _normalize_unit(val: float, raw_text: str) -> tuple:
    """Normalize LLM metric value to 0-1 range and detect unit."""
    lower = raw_text.lower()
    if "%" in raw_text or "percent" in lower:
        return "percent", val / 100.0
    if val > 1.0 and any(m in lower for m in ["accuracy", "acc", "f1", "bleu", "rouge", "win rate", "success rate"]):
        # Heuristic: values > 1 for accuracy-like metrics are likely percentages
        return "percent", val / 100.0
    return "ratio", val


# Standalone sample size patterns
N_PATTERNS = [
    re.compile(r'\b(?:N|n)\s*=\s*(\d+)', re.IGNORECASE),
    re.compile(r'\b(\d+)\s+(?:participants|subjects|patients|samples|queries|tasks|instances|images)\b', re.IGNORECASE),
]


def extract_standalone_n(text: str, location: str = "text") -> List[Dict[str, Any]]:
    """Extract standalone sample size mentions not embedded in descriptives/NHST."""
    results = []
    for pattern in N_PATTERNS:
        for match in pattern.finditer(text):
            try:
                val = int(match.group(1))
                results.append({
                    "entity_type": "sample_size",
                    "section": location,
                    "value": val,
                    "raw_text": match.group(0).strip(),
                })
            except (ValueError, IndexError):
                continue
    return results


def link_proportion_with_descriptive(
    proportions: List[Dict], descriptives: List[Dict], max_distance: int = 200
) -> List[Dict]:
    """
    P0: If a proportion and a descriptive (M, SD, n) are nearby with same n
    and mean ≈ proportion, inject descriptive.sd as proportion.reported_sd.
    """
    if not proportions or not descriptives:
        return proportions

    linked = []
    for prop in proportions:
        prop_text = prop.get("raw_text", "")
        prop_n = prop.get("total_n")
        prop_mean = prop.get("proportion")

        best_sd = prop.get("reported_sd")  # keep existing if any
        for desc in descriptives:
            desc_n = desc.get("n")
            desc_mean = desc.get("mean")
            desc_sd = desc.get("sd")
            if desc_n is None or desc_sd is None or desc_mean is None:
                continue
            # Same n and mean ≈ proportion (within 5% relative)
            if prop_n and desc_n == prop_n:
                rel_diff = abs(desc_mean - prop_mean) / max(prop_mean, 0.001)
                if rel_diff < 0.05:
                    best_sd = desc_sd
                    break

        prop["reported_sd"] = best_sd
        linked.append(prop)
    return linked


def extract_proportions(text: str, location: str = "text") -> List[ProportionResult]:
    """Extract reported proportions WITHOUT computing SD ourselves."""
    results = []
    for pattern in PROPORTION_PATTERNS:
        for match in pattern.finditer(text):
            try:
                groups = match.groups()
                raw = match.group(0)

                # Pattern 1: 45% (18/40)
                if "%" in raw and groups[1] and groups[2]:
                    pct = float(groups[0]) / 100
                    count = int(groups[1])
                    total = int(groups[2])
                # Pattern 2: 18/40 (45%)
                elif "/" in raw and groups[0] and groups[1]:
                    count = int(groups[0])
                    total = int(groups[1])
                    pct = float(groups[2]) / 100
                else:
                    continue

                # P0: Look for nearby M, SD, n in same sentence/clause
                reported_sd = None
                clause_end = min(len(text), match.end() + 120)
                nearby = text[match.start():clause_end]
                sd_pat = re.search(r"(?:SD|sd)\s*[=:]\s*" + NUM, nearby)
                if sd_pat:
                    reported_sd = float(re.search(NUM, sd_pat.group(0)).group(0))

                results.append(ProportionResult(
                    location=location,
                    raw_text=raw.strip(),
                    proportion=pct,
                    count=count,
                    total_n=total,
                    reported_sd=reported_sd,
                ))
            except (ValueError, TypeError):
                continue
    return results


def extract_llm_metrics(text: str, location: str = "text") -> List[LLMMetricResult]:
    results = []
    for pattern in LLM_PATTERNS:
        for match in pattern.finditer(text):
            try:
                groups = match.groups()
                raw = match.group(0)
                lower = raw.lower()
                n_groups = len(groups)

                # pass@k: group 0=k, group 1=value
                if "pass" in lower and "@" in raw:
                    if n_groups >= 2:
                        k = int(groups[0])
                        val = float(groups[1])
                        unit, norm = _normalize_unit(val, raw)
                        results.append(LLMMetricResult(
                            location=location, raw_text=raw.strip(),
                            metric=f"pass@{k}", value=val,
                            value_normalized=norm, unit=unit,
                        ))
                # std over seeds: group 0=seed count
                elif "seed" in lower and n_groups >= 1:
                    seeds = int(groups[0])
                    results.append(LLMMetricResult(
                        location=location, raw_text=raw.strip(),
                        metric="std_over_seeds", value=seeds,
                        seeds=seeds,
                    ))
                # All others: group 0=value
                elif n_groups >= 1 and groups[0]:
                    val = float(groups[0])
                    metric_name = "metric"
                    if "win" in lower:
                        metric_name = "win_rate"
                    elif "success" in lower:
                        metric_name = "success_rate"
                    elif "em" in lower or "exact match" in lower:
                        metric_name = "exact_match"
                    elif "accuracy" in lower or "acc" in lower:
                        metric_name = "accuracy"
                    elif "f1" in lower:
                        metric_name = "f1"
                    elif "bleu" in lower:
                        metric_name = "bleu"
                    elif "rouge" in lower:
                        metric_name = "rouge"
                    elif "auroc" in lower or "auc" in lower:
                        metric_name = "auroc"
                    elif "auprc" in lower:
                        metric_name = "auprc"
                    elif "latency" in lower:
                        metric_name = "latency"
                    elif "token" in lower:
                        metric_name = "tokens"
                    elif "cost" in lower:
                        metric_name = "cost"
                    elif "api call" in lower or "tool call" in lower:
                        metric_name = "tool_calls"
                    unit, norm = _normalize_unit(val, raw)
                    results.append(LLMMetricResult(
                        location=location, raw_text=raw.strip(),
                        metric=metric_name, value=val,
                        value_normalized=norm, unit=unit,
                    ))
                # mean ± std: group 0=mean, group 1=std
                elif n_groups >= 2 and groups[0] and groups[1]:
                    results.append(LLMMetricResult(
                        location=location, raw_text=raw.strip(),
                        metric="mean_std", value=float(groups[0]),
                        std=float(groups[1]),
                    ))
            except (ValueError, TypeError, IndexError):
                continue
    return results


# ---------------------------------------------------------------------------
# Document extraction
# ---------------------------------------------------------------------------

def extract_from_docx(filepath: str) -> Dict[str, Any]:
    if not HAS_DOCX:
        return {"error": "python-docx not installed. pip install python-docx"}
    doc = docx.Document(filepath)
    sections = {}
    current_section = "document_start"
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        lower = text.lower()
        if any(h in lower for h in ["abstract", "introduction", "methods", "results",
                                      "discussion", "conclusion", "references",
                                      "supplementary", "appendix"]):
            current_section = text.replace(":", "").strip()
        if current_section not in sections:
            sections[current_section] = []
        sections[current_section].append(text)

    for i, table in enumerate(doc.tables):
        table_texts = []
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells]
            table_texts.append("\t".join(row_texts))
        loc = f"Table {i+1}"
        if "tables" not in sections:
            sections["tables"] = {}
        sections["tables"][loc] = "\n".join(table_texts)
    return sections


def extract_from_pdf(filepath: str) -> Dict[str, Any]:
    try:
        if HAS_PDFMINER:
            text = pdf_extract_text(filepath)
        elif HAS_FITZ:
            doc = fitz.open(filepath)
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
        else:
            return {"error": "No PDF library. pip install pdfminer.six or PyMuPDF."}

        sections = {"document_start": []}
        current = "document_start"
        markers = ["abstract", "introduction", "methods", "results",
                   "discussion", "conclusion", "references", "supplementary", "appendix"]
        for line in text.split("\n"):
            line = line.strip()
            if not line:
                continue
            lower = line.lower()
            for m in markers:
                if lower.startswith(m) or lower == m:
                    current = m
                    if current not in sections:
                        sections[current] = []
                    break
            if current not in sections:
                sections[current] = []
            sections[current].append(line)
        return sections
    except Exception as e:
        return {"error": str(e)}


def extract_all_stats(text: str, location: str = "text") -> Dict[str, Any]:
    nhst = [asdict(r) for r in extract_nhst(text, location)]
    descriptives = [asdict(r) for r in extract_descriptives(text, location)]
    proportions = [asdict(r) for r in extract_proportions(text, location)]

    # P0: Link nearby proportion + M/SD/n for DEBIT
    prop_dicts = [dict(p) for p in proportions]
    desc_dicts = [dict(d) for d in descriptives]
    linked_props = link_proportion_with_descriptive(prop_dicts, desc_dicts)

    return {
        "nhst": nhst,
        "descriptives": descriptives,
        "proportions": linked_props,
        "standalone_n": extract_standalone_n(text, location),
        "llm_metrics": [asdict(r) for r in extract_llm_metrics(text, location)],
    }


def main():
    parser = argparse.ArgumentParser(description="Extract statistics from manuscript")
    parser.add_argument("--input", "-i", help="Input file (.docx or .pdf)")
    parser.add_argument("--text", "-t", help="Direct text input")
    parser.add_argument("--output", "-o", default="stats_extracted.json", help="Output JSON")
    args = parser.parse_args()

    if not args.input and not args.text:
        parser.print_help()
        sys.exit(1)

    all_results = {}

    if args.input:
        filepath = Path(args.input)
        if filepath.suffix.lower() == ".docx":
            sections = extract_from_docx(str(filepath))
        elif filepath.suffix.lower() == ".pdf":
            sections = extract_from_pdf(str(filepath))
        else:
            print(f"Unsupported: {filepath.suffix}")
            sys.exit(1)

        if "error" in sections:
            print(f"Error: {sections['error']}")
            sys.exit(1)

        for section_name, content in sections.items():
            if isinstance(content, dict):
                for table_name, table_text in content.items():
                    loc = f"{section_name} - {table_name}"
                    all_results[loc] = extract_all_stats(table_text, loc)
            elif isinstance(content, list):
                all_results[section_name] = extract_all_stats("\n".join(content), section_name)
            else:
                all_results[section_name] = extract_all_stats(str(content), section_name)

    elif args.text:
        all_results["input_text"] = extract_all_stats(args.text, "input_text")

    with open(args.output, "w", encoding="utf-8") as f:
        json.dump(all_results, f, indent=2, ensure_ascii=False)

    total_nhst = sum(len(r.get("nhst", [])) for r in all_results.values())
    total_desc = sum(len(r.get("descriptives", [])) for r in all_results.values())
    total_prop = sum(len(r.get("proportions", [])) for r in all_results.values())
    total_n = sum(len(r.get("standalone_n", [])) for r in all_results.values())
    total_llm = sum(len(r.get("llm_metrics", [])) for r in all_results.values())

    print(f"Extraction complete -> {args.output}")
    print(f"  NHST: {total_nhst}, Descriptives: {total_desc}, Proportions: {total_prop}, Standalone N: {total_n}, LLM metrics: {total_llm}")


if __name__ == "__main__":
    main()
